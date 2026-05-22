from __future__ import annotations

import json
import mimetypes
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from pathlib import Path

from docx import Document


BASE_URL = "http://localhost:8080"
OLLAMA_URL = "http://localhost:11434"
ROOT = Path(r"C:\CurserWorkSpace\docuagent-local")
OUT = ROOT / "storage" / "phase16-qa"
OUT.mkdir(parents=True, exist_ok=True)


TEMPLATES = [
    {
        "kind": "education",
        "path": ROOT / "storage" / "phase16-templates" / "교육계획안_realistic.docx",
        "tab_name": "QA 교육계획안",
        "description": "표 기반 교육계획안 양식",
        "base_prompt": "교육계획안 문체로 간결하고 전문적으로 정리한다.",
        "task_title": "3월 2주차 봄꽃 관찰 교육계획안",
        "context": (
            "활동 목표: 유아가 봄꽃을 관찰하며 계절 변화를 이해한다. "
            "준비물: 색종이, 가위, 풀, 봄꽃 사진 카드. "
            "활동 내용: 교실 창가에 놓인 꽃을 관찰하고 색종이로 꽃을 만들어 보았다. "
            "유아 반응: 일부 유아는 꽃 이름을 헷갈려 했지만 색과 향기에 관심을 보였다. "
            "특이사항: 가위 사용이 어려운 유아에게 개별 지원이 필요했다."
        ),
        "expected_keys": ["activityGoal", "materials", "activityContent", "childReaction", "notes"],
    },
    {
        "kind": "observation",
        "path": ROOT / "storage" / "phase16-templates" / "관찰일지_realistic.docx",
        "tab_name": "QA 관찰일지",
        "description": "병합 셀 제목이 포함된 관찰일지 양식",
        "base_prompt": "관찰 기록은 사실과 해석을 구분해 작성한다.",
        "task_title": "김민준 놀이 관찰 기록",
        "context": (
            "관찰: 김민준은 블록으로 다리를 만들며 친구에게 파란 블록을 건네 달라고 말했다. "
            "분석: 또래와 역할을 나누어 구성 놀이를 이어 가는 모습이 관찰되었다. "
            "지원계획: 다음 활동에서 긴 블록과 연결 블록을 충분히 제공한다. "
            "특이사항: 놀이 중 블록이 무너졌을 때 잠시 속상해했으나 다시 시도했다."
        ),
        "expected_keys": ["observation", "analysis", "supportPlan", "notes"],
    },
    {
        "kind": "counseling",
        "path": ROOT / "storage" / "phase16-templates" / "상담일지_realistic.docx",
        "tab_name": "QA 상담일지",
        "description": "상담 항목과 추후 계획이 있는 상담일지 양식",
        "base_prompt": "상담일지는 보호자 의견과 교사 의견을 분리해 정리한다.",
        "task_title": "5월 보호자 상담일지",
        "context": (
            "상담 내용: 보호자는 최근 등원 시 아이가 분리 불안을 보인다고 말했다. "
            "보호자 의견: 가정에서는 잠들기 전 어린이집 이야기를 자주 묻고 있다고 했다. "
            "교사 의견: 교실에서는 등원 후 10분 정도 지나면 또래와 안정적으로 놀이한다. "
            "추후 계획: 등원 직후 담임교사가 짧은 인사 루틴을 반복하고 보호자에게 적응 상황을 공유한다."
        ),
        "expected_keys": ["consultationContent", "parentOpinion", "teacherOpinion", "followUpPlan"],
    },
]


def request_json(method: str, path: str, body=None, timeout=60):
    data = None
    headers = {}
    if body is not None:
        data = json.dumps(body, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json; charset=utf-8"
    req = urllib.request.Request(BASE_URL + path, data=data, method=method, headers=headers)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        raw = resp.read()
        return json.loads(raw.decode("utf-8")) if raw else None


def multipart_post(path: str, fields: dict[str, str], file_field: str, file_path: Path, timeout=60):
    boundary = f"----docuagent-phase16-{int(time.time() * 1000)}"
    parts: list[bytes] = []
    for name, value in fields.items():
        parts.append(f"--{boundary}\r\n".encode("utf-8"))
        parts.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"))
        parts.append(value.encode("utf-8"))
        parts.append(b"\r\n")
    filename = file_path.name
    content_type = mimetypes.guess_type(filename)[0] or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    parts.append(f"--{boundary}\r\n".encode("utf-8"))
    parts.append(
        (
            f'Content-Disposition: form-data; name="{file_field}"; filename="{filename}"\r\n'
            f"Content-Type: {content_type}\r\n\r\n"
        ).encode("utf-8")
    )
    parts.append(file_path.read_bytes())
    parts.append(b"\r\n")
    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    data = b"".join(parts)
    req = urllib.request.Request(
        BASE_URL + path,
        data=data,
        method="POST",
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}", "Content-Length": str(len(data))},
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return json.loads(resp.read().decode("utf-8"))


def download(path: str, output_path: Path, timeout=90):
    req = urllib.request.Request(BASE_URL + path, method="GET")
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        content_type = resp.headers.get("Content-Type")
        content_disposition = resp.headers.get("Content-Disposition")
        output_path.write_bytes(resp.read())
        return {"contentType": content_type, "contentDisposition": content_disposition, "bytes": output_path.stat().st_size}


def ollama_tags():
    with urllib.request.urlopen(OLLAMA_URL + "/api/tags", timeout=10) as resp:
        return json.loads(resp.read().decode("utf-8"))


def structured_from_document(document_response):
    structured = document_response.get("structuredContent")
    if structured:
        return structured
    return json.loads(document_response["generatedContent"])


def inspect_docx(path: Path):
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_cells = []
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    table_cells.append(
                        {
                            "table": table_index,
                            "row": row_index,
                            "cell": cell_index,
                            "text": text,
                        }
                    )
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    return {
        "paragraphs": paragraphs,
        "tableCells": table_cells,
        "hasPlaceholderSyntax": "{{" in xml or "}}" in xml,
        "xmlChars": len(xml),
    }


def value_presence(structured: dict[str, str], inspection: dict):
    haystack = "\n".join(inspection["paragraphs"] + [cell["text"] for cell in inspection["tableCells"]])
    result = {}
    for key, value in structured.items():
        normalized = (value or "").strip()
        if not normalized or normalized in {"추가 입력 필요", "미기재"}:
            result[key] = "NO_VALUE"
            continue
        tokens = [token for token in re.split(r"\s+", normalized) if len(token) >= 2]
        matched = sum(1 for token in tokens[:8] if token in haystack)
        result[key] = "FOUND" if normalized in haystack or matched >= max(1, min(3, len(tokens))) else "NOT_FOUND"
    return result


def run_one(template):
    report = {
        "kind": template["kind"],
        "templatePath": str(template["path"]),
        "steps": {},
    }

    tab = multipart_post(
        "/api/tabs",
        {
            "name": template["tab_name"],
            "description": template["description"],
            "basePrompt": template["base_prompt"],
        },
        "templateFile",
        template["path"],
    )
    tab_id = tab["id"]
    report["tabId"] = tab_id
    report["steps"]["upload"] = {"ok": True, "originalFileName": tab.get("originalFileName")}

    analysis = request_json("POST", f"/api/tabs/{tab_id}/analyze-template", timeout=90)
    labels = [label["text"] for label in analysis.get("labels", [])]
    mappings = request_json("GET", f"/api/tabs/{tab_id}/template-mappings")
    mapping_by_key = {m.get("semanticFieldKey"): m for m in mappings if m.get("semanticFieldKey")}
    expected_mapping_missing = [key for key in template["expected_keys"] if key not in mapping_by_key]
    report["steps"]["analysis"] = {
        "ok": True,
        "blockCount": len(analysis.get("blocks", [])),
        "labelCount": len(labels),
        "labels": labels,
        "mappings": mappings,
        "expectedMappingMissing": expected_mapping_missing,
    }

    task = request_json(
        "POST",
        f"/api/tabs/{tab_id}/tasks",
        {"title": template["task_title"], "userContext": template["context"]},
    )
    task_id = task["id"]
    report["taskId"] = task_id
    report["steps"]["task"] = {"ok": True, "title": task["title"]}

    generated = request_json("POST", f"/api/tasks/{task_id}/generate-draft", timeout=240)
    structured = structured_from_document(generated)
    report["documentId"] = generated["id"]
    report["steps"]["generation"] = {
        "ok": True,
        "documentId": generated["id"],
        "structured": structured,
        "missingValues": [k for k, v in structured.items() if not v or v.strip() in {"추가 입력 필요", "미기재"}],
    }

    saved = request_json(
        "PUT",
        f"/api/documents/{generated['id']}",
        {"generatedContent": json.dumps(structured, ensure_ascii=False)},
    )
    saved_structured = structured_from_document(saved)
    report["steps"]["saveDraft"] = {"ok": True, "updatedAt": saved.get("updatedAt")}

    preview = request_json("GET", f"/api/tasks/{task_id}/placement-preview")
    plan = request_json("GET", f"/api/tasks/{task_id}/write-plan")
    report["steps"]["placementPreview"] = {
        "ok": True,
        "rows": preview,
        "statusCounts": count_by(preview, "status"),
    }
    report["steps"]["writePlan"] = {
        "ok": True,
        "operations": plan,
        "statusCounts": count_by(plan, "status"),
    }

    download_path = OUT / f"{template['kind']}-{task_id}-downloaded.docx"
    download_info = download(f"/api/tasks/{task_id}/download-docx", download_path)
    summary = request_json("GET", f"/api/tasks/{task_id}/reconstruction-summary")
    inspection = inspect_docx(download_path)
    report["steps"]["download"] = {
        "ok": True,
        "path": str(download_path),
        **download_info,
    }
    report["steps"]["summary"] = summary
    report["steps"]["inspection"] = {
        "ok": True,
        "hasPlaceholderSyntax": inspection["hasPlaceholderSyntax"],
        "valuePresence": value_presence(saved_structured, inspection),
        "paragraphCount": len(inspection["paragraphs"]),
        "tableCellCount": len(inspection["tableCells"]),
        "sampleCells": inspection["tableCells"][:16],
    }

    return report


def count_by(rows, key):
    counts = {}
    for row in rows:
        value = row.get(key)
        counts[value] = counts.get(value, 0) + 1
    return counts


def main():
    final = {
        "settings": None,
        "ollama": None,
        "templates": [],
        "errors": [],
    }
    try:
        final["settings"] = request_json("GET", "/api/settings")
    except Exception as exc:
        final["errors"].append({"stage": "settings", "error": repr(exc)})
    try:
        tags = ollama_tags()
        final["ollama"] = {
            "available": True,
            "models": [model.get("name") for model in tags.get("models", [])],
        }
    except Exception as exc:
        final["ollama"] = {"available": False, "error": repr(exc)}

    for template in TEMPLATES:
        try:
            final["templates"].append(run_one(template))
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="replace")
            final["templates"].append(
                {
                    "kind": template["kind"],
                    "templatePath": str(template["path"]),
                    "error": f"HTTP {exc.code}",
                    "body": body,
                }
            )
        except Exception as exc:
            final["templates"].append(
                {
                    "kind": template["kind"],
                    "templatePath": str(template["path"]),
                    "error": repr(exc),
                }
            )

    report_path = OUT / "phase16-qa-report.json"
    report_path.write_text(json.dumps(final, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(final, ensure_ascii=True, indent=2))
    print(f"REPORT={report_path}")


if __name__ == "__main__":
    sys.exit(main())
