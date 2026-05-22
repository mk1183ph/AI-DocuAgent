from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


OUT = Path(r"C:\CurserWorkSpace\docuagent-local\storage\phase16-templates")
OUT.mkdir(parents=True, exist_ok=True)


def set_font(run, size=10, bold=False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold


def add_title(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, 16, True)


def add_label_table(document, rows, widths=(1.8, 4.7)):
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, (label, sample) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = ""
        label_run = cells[0].paragraphs[0].add_run(label)
        set_font(label_run, 10, True)
        cells[1].text = sample
        if sample:
            for run in cells[1].paragraphs[0].runs:
                set_font(run, 10)
    return table


def save(document, name):
    path = OUT / name
    document.save(path)
    print(path)


def education_plan():
    document = Document()
    add_title(document, "교육계획안")
    info = document.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    labels = ["반명", "", "기간", ""]
    labels2 = ["주제", "", "담당교사", ""]
    for row, values in zip(info.rows, [labels, labels2]):
        for cell, value in zip(row.cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, 9, bool(value))
    document.add_paragraph()
    add_label_table(document, [
        ("활동 목표", ""),
        ("준비물", ""),
        ("활동 내용", ""),
        ("유아 반응", ""),
        ("특이사항", ""),
    ])
    save(document, "교육계획안_realistic.docx")


def observation_log():
    document = Document()
    add_title(document, "유아 관찰일지")
    meta = document.add_table(rows=1, cols=4)
    meta.style = "Table Grid"
    for cell, text in zip(meta.rows[0].cells, ["유아명", "", "관찰일", ""]):
        cell.text = text
    document.add_paragraph()
    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[0].cells[0].text = "관찰 기록"
    rows = [
        ("관찰", ""),
        ("분석", ""),
        ("지원계획", ""),
        ("특이사항", ""),
    ]
    for idx, (label, sample) in enumerate(rows, start=1):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = sample
        for cell in table.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    save(document, "관찰일지_realistic.docx")


def counseling_log():
    document = Document()
    add_title(document, "상담일지")
    table = document.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[0].cells[0].text = "보호자 상담 기록"
    rows = [
        ("상담일", ""),
        ("상담 대상", ""),
        ("상담 내용", ""),
        ("보호자 의견", ""),
        ("교사 의견", ""),
        ("추후 계획", ""),
    ]
    for idx, (label, sample) in enumerate(rows, start=1):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = sample
    save(document, "상담일지_realistic.docx")


education_plan()
observation_log()
counseling_log()
